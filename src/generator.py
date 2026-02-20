import io
from typing import List
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(images: List[Image.Image], output_path: str):
    """
    โมดูลการประกอบและสร้างเอกสารผลลัพธ์ (DOCX):
    นำภาพคู่ขนานทั้งหมดมาจัดเรียงลงในไฟล์ Word โดยบังคับหน้ากระดาษเป็นแนวนอน
    """
    print("กำลังสร้างไฟล์เอกสาร DOCX...")
    docx_document = Document()

    # ปรับแต่ง Section แรกให้แสดงผลแบบแนวนอน (Landscape Orientation) เพื่อรองรับภาพกว้าง
    first_section = docx_document.sections[0]
    original_width, original_height = (
        first_section.page_width,
        first_section.page_height,
    )
    first_section.orientation = 1  # รหัส 1 แทนความหมายของ Landscape
    first_section.page_width = original_height
    first_section.page_height = original_width

    # บีบระยะขอบกระดาษ (Margins) ให้แคบลง (0.5 นิ้ว) เพื่อเพิ่มพื้นที่วางรูปภาพให้กว้างขึ้น
    margin_size = Inches(0.5)
    first_section.left_margin = margin_size
    first_section.right_margin = margin_size
    first_section.top_margin = margin_size
    first_section.bottom_margin = margin_size

    for idx, img in enumerate(images):
        # ใช้ I/O stream แปลงภาพพักไว้ในหน่วยความจำ (ไม่ต้องบันทึกลงฮาร์ดดิสก์ให้รก)
        virtual_image_stream = io.BytesIO()
        img.save(virtual_image_stream, format="PNG")
        virtual_image_stream.seek(0)

        # สร้างย่อหน้าใหม่ จัดกึ่งกลาง และบรรจุภาพลงไป
        paragraph_container = docx_document.add_paragraph()
        paragraph_container.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ปรับแก้พฤติกรรม Paragraph ดันบรรทัด (Spacing Overflow) ที่ทำให้เกิดหน้าว่าง (Empty Page Bug)
        paragraph_format = paragraph_container.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        run_element = paragraph_container.add_run()

        # กำหนดความกว้างภาพประมาณ 9.5 นิ้ว เพื่อให้พอดีกับกระดาษแนวนอน A4/Letter
        run_element.add_picture(virtual_image_stream, width=Inches(9.5))

        # คำสั่งขึ้นหน้าใหม่ (Page Break) ยกเว้นหน้าสุดท้าย
        if idx < len(images) - 1:
            docx_document.add_page_break()

    docx_document.save(output_path)
    print(f"[สำเร็จ] รายงานผลลัพธ์ DOCX ถูกบันทึกที่: {output_path}")


def generate_pdf(images: List[Image.Image], output_path: str):
    """
    โมดูลการประกอบและสร้างเอกสารผลลัพธ์ (PDF):
    อาศัยความสามารถของ Pillow ในการบันทึกเอกสารหลายหน้า (Multi-page PDF save)
    """
    print("กำลังสร้างไฟล์เอกสาร PDF...")
    if not images:
        print("ไม่มีภาพสำหรับสร้าง PDF")
        return

    cover_page = images[0]
    subsequent_pages = images[1:]

    # พารามิเตอร์ save_all=True และ append_images สั่งให้รวมพิกเซลทั้งหมดเป็น PDF เดียว
    cover_page.save(
        output_path,
        "PDF",
        resolution=150.0,
        save_all=True,
        append_images=subsequent_pages,
    )
    print(f"[สำเร็จ] รายงานผลลัพธ์ PDF ถูกบันทึกที่: {output_path}")
